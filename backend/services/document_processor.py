import os
import uuid
import time
from typing import List, Dict, Any
from pathlib import Path
import PyPDF2
from docx import Document
import chromadb
from chromadb.config import Settings

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument

from utils.config import config
from utils.text_processing import text_processor
from utils.logger import log_document_processing, log_error
from models.document import DocumentType, DocumentStatus
from services.embedding_service import embedding_service


class DocumentProcessor:
    def __init__(self):
        self.chroma_client = chromadb.PersistentClient(
            path=config.chroma_db_path,
            settings=Settings(anonymized_telemetry=False)
        )

        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.chunk_size,
            chunk_overlap=config.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )

        # Get or create collection with optimized HNSW parameters
        self.collection = self.chroma_client.get_or_create_collection(
            name="ragify_documents",
            metadata={
                "hnsw:space": "cosine",
                "hnsw:construction_ef": 64,  # More conservative construction ef
                "hnsw:M": 8                   # More conservative M value
            }
        )

    def extract_text_from_pdf(self, file_path: str) -> tuple[str, list[dict]]:
        """Extract text from PDF file with page information"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                page_info = []

                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    text += page_text + "\n"

                    # Store page information
                    page_info.append({
                        "page_number": page_num,
                        "text": page_text,
                        # -1 for the newline
                        "start_char": len(text) - len(page_text) - 1,
                        "end_char": len(text) - 1
                    })

                return text.strip(), page_info
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")

    def extract_text_from_docx(self, file_path: str) -> tuple[str, list[dict]]:
        """Extract text from DOCX file with page information"""
        try:
            doc = Document(file_path)
            text = ""
            page_info = []

            # For DOCX, we'll treat each paragraph as a "page" for simplicity
            # In a real implementation, you might want to use a library that can detect page breaks
            for para_num, paragraph in enumerate(doc.paragraphs, 1):
                para_text = paragraph.text
                if para_text.strip():  # Only add non-empty paragraphs
                    text += para_text + "\n"
                    page_info.append({
                        "page_number": para_num,
                        "text": para_text,
                        "start_char": len(text) - len(para_text) - 1,
                        "end_char": len(text) - 1
                    })

            return text.strip(), page_info
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")

    def extract_text_from_txt(self, file_path: str) -> tuple[str, list[dict]]:
        """Extract text from TXT file with page information"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read().strip()
                # For TXT files, we'll treat each paragraph as a "page"
                paragraphs = content.split('\n\n')
                page_info = []

                for para_num, paragraph in enumerate(paragraphs, 1):
                    if paragraph.strip():
                        page_info.append({
                            "page_number": para_num,
                            "text": paragraph,
                            "start_char": content.find(paragraph),
                            "end_char": content.find(paragraph) + len(paragraph)
                        })

                return content, page_info
        except Exception as e:
            raise Exception(f"Error reading text file: {str(e)}")

    def extract_text(self, file_path: str, document_type: DocumentType) -> tuple[str, list[dict]]:
        """Extract text based on document type with page information"""
        if document_type == DocumentType.PDF:
            return self.extract_text_from_pdf(file_path)
        elif document_type == DocumentType.DOCX:
            return self.extract_text_from_docx(file_path)
        elif document_type == DocumentType.TEXT:
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported document type: {document_type}")

    def find_page_number(self, chunk_start: int, page_info: list[dict]) -> int:
        """Find the page number for a given character position"""
        for page in page_info:
            if page["start_char"] <= chunk_start <= page["end_char"]:
                return page["page_number"]
        return 1  # Default to page 1 if not found

    def process_document(self, file_path: str, filename: str, document_type: DocumentType) -> Dict[str, Any]:
        """Process a document: extract text, chunk, and create embeddings"""
        start_time = time.time()
        extraction_start = None
        embedding_start = None
        chromadb_start = None

        try:
            print(f"Starting to process document: {filename} ({file_path})")

            # Check if document already exists in ChromaDB with more specific criteria
            # Use both filename and file path to avoid false positives
            existing_docs = self.collection.get(
                where={"source": filename}
            )

            # Only skip if we have existing docs AND the file path matches
            if existing_docs['ids'] and len(existing_docs['ids']) > 0:
                # Check if the file path is actually the same (more thorough check)
                # For now, we'll process anyway to ensure fresh embeddings
                print(
                    f"Document {filename} found in ChromaDB, but reprocessing for fresh embeddings...")

                # Optionally, you could delete existing embeddings first:
                # self.collection.delete(where={"source": filename})

            print(f"Extracting text from {filename}...")
            extraction_start = time.time()
            # Extract text
            raw_text, page_info = self.extract_text(file_path, document_type)
            extraction_time = time.time() - extraction_start
            print(
                f"Extracted {len(raw_text)} characters from {filename} in {extraction_time:.2f}s")

            # Clean text
            cleaned_text = text_processor.clean_text(raw_text)
            print(f"Cleaned text: {len(cleaned_text)} characters")

            # Split into chunks
            chunks = self.text_splitter.split_text(cleaned_text)
            print(f"Split into {len(chunks)} chunks")

            # Create documents for embedding
            documents = []
            current_pos = 0

            for i, chunk in enumerate(chunks):
                # Find which page this chunk belongs to
                chunk_start = current_pos
                page_number = self.find_page_number(chunk_start, page_info)

                doc = LangchainDocument(
                    page_content=chunk,
                    metadata={
                        "source": filename,
                        "chunk_index": i,
                        "document_type": document_type.value,
                        "total_chunks": len(chunks),
                        "page_number": page_number,
                        "file_path": file_path  # Add file path to metadata
                    }
                )
                documents.append(doc)
                current_pos += len(chunk) + 1  # +1 for the separator

            print(f"Generating embeddings for {len(documents)} chunks...")
            embedding_start = time.time()
            # Generate embeddings using GCP Gemini
            embeddings_list = embedding_service.embed_documents(
                [doc.page_content for doc in documents])
            embedding_time = time.time() - embedding_start
            print(f"Generated embeddings in {embedding_time:.2f}s")

            # Prepare data for ChromaDB
            ids = [f"{filename}_{i}" for i in range(len(documents))]
            texts = [doc.page_content for doc in documents]
            metadatas = [doc.metadata for doc in documents]

            # Remove existing embeddings for this document first
            self.collection.delete(where={"source": filename})

            print(f"Adding {len(documents)} chunks to ChromaDB...")
            chromadb_start = time.time()
            # Add to ChromaDB
            self.collection.add(
                embeddings=embeddings_list,
                documents=texts,
                metadatas=metadatas,
                ids=ids
            )
            chromadb_time = time.time() - chromadb_start
            print(f"Added to ChromaDB in {chromadb_time:.2f}s")

            total_time = time.time() - start_time
            result = {
                "chunks_count": len(chunks),
                "total_tokens": len(cleaned_text.split()),
                "keywords": text_processor.extract_keywords(cleaned_text, max_keywords=10),
                "timing": {
                    "extraction_time": extraction_time,
                    "embedding_time": embedding_time,
                    "chromadb_time": chromadb_time,
                    "total_time": total_time
                }
            }

            print(
                f"Successfully processed {filename}: {result['chunks_count']} chunks, {result['total_tokens']} tokens in {total_time:.2f}s")
            return result

        except Exception as e:
            total_time = time.time() - start_time
            print(
                f"Error processing document {filename}: {str(e)} (took {total_time:.2f}s)")
            log_error(e, f"Document processing failed for {filename}")
            raise Exception(f"Error processing document: {str(e)}")

    def has_sufficient_data(self) -> bool:
        """Check if the collection has enough data for searching"""
        try:
            count = self.collection.count()
            return count >= 3  # Need at least 3 chunks for meaningful search
        except Exception:
            return False

    def search_documents(self, query: str, document_ids: List[str] | None = None, limit: int = 10) -> List[Dict[str, Any]]:
        """Search documents using semantic similarity"""
        try:
            # Check if we have enough data to search
            if not self.has_sufficient_data():
                raise Exception(
                    "Insufficient data in the search index. Please upload more documents and try again.")

            # Generate query embedding
            query_embedding = embedding_service.embed_query(query)

            # Prepare where clause for filtering by document
            where_clause = None
            if document_ids:
                where_clause = {"source": {"$in": document_ids}}
                print(f"DEBUG: Using where clause: {where_clause}")
                print(f"DEBUG: Document IDs for filtering: {document_ids}")
            else:
                print("DEBUG: No document filtering - searching all documents")

            # Try different search parameters if the first attempt fails
            search_attempts = [
                {"n_results": limit},  # Default attempt
                {"n_results": min(limit, 5)},  # Reduced results
                {"n_results": min(limit, 3)},  # Further reduced results
            ]

            last_error = Exception("All search attempts failed")
            for attempt, params in enumerate(search_attempts):
                try:
                    # Search in ChromaDB
                    results = self.collection.query(
                        query_embeddings=[query_embedding],
                        n_results=params["n_results"],
                        where=where_clause
                    )

                    # Format results
                    formatted_results = []
                    if results['documents'] and results['documents'][0]:
                        for i, (doc, metadata, distance) in enumerate(zip(
                            results['documents'][0],
                            results['metadatas'][0],
                            results['distances'][0]
                        )):
                            formatted_results.append({
                                "id": results['ids'][0][i],
                                "content": doc,
                                "metadata": metadata,
                                "score": 1 - distance,  # Convert distance to similarity score
                                "source": metadata.get("source", "Unknown"),
                                "page_number": metadata.get("page_number", 1)
                            })

                    return formatted_results

                except Exception as e:
                    last_error = e
                    error_msg = str(e)

                    # If it's the HNSW error, try the next attempt
                    if "Cannot return the results in a contigious 2D array" in error_msg:
                        print(
                            f"Search attempt {attempt + 1} failed with HNSW error, trying with different parameters...")
                        continue
                    else:
                        # If it's a different error, don't retry
                        break

            # If all attempts failed, raise the last error
            raise last_error

        except Exception as e:
            error_msg = str(e)
            if "Cannot return the results in a contigious 2D array" in error_msg:
                raise Exception(
                    "Search failed due to insufficient data in the index. Please try uploading more documents or try a different search query.")
            else:
                raise Exception(f"Error searching documents: {error_msg}")

    def get_document_chunks(self, filename: str) -> List[Dict[str, Any]]:
        """Get all chunks for a specific document"""
        try:
            results = self.collection.get(
                where={"source": filename}
            )

            chunks = []
            for i, (doc, metadata) in enumerate(zip(results['documents'], results['metadatas'])):
                chunks.append({
                    "id": results['ids'][i],
                    "content": doc,
                    "metadata": metadata
                })

            return chunks

        except Exception as e:
            raise Exception(f"Error retrieving document chunks: {str(e)}")

    def delete_document(self, filename: str) -> bool:
        """Delete all chunks for a specific document"""
        try:
            self.collection.delete(
                where={"source": filename}
            )
            return True
        except Exception as e:
            raise Exception(f"Error deleting document: {str(e)}")

    def reset_collection(self):
        """Reset the ChromaDB collection with optimized parameters"""
        try:
            # Delete the existing collection
            self.chroma_client.delete_collection("ragify_documents")
            print("Deleted existing ChromaDB collection")

            # Recreate with optimized parameters
            self.collection = self.chroma_client.create_collection(
                name="ragify_documents",
                metadata={
                    "hnsw:space": "cosine",
                    "hnsw:construction_ef": 64,  # More conservative construction ef
                    "hnsw:M": 8                   # More conservative M value
                }
            )
            print("Recreated ChromaDB collection with optimized parameters")
            return True
        except Exception as e:
            print(f"Error resetting collection: {str(e)}")
            return False

    def get_collection_stats(self) -> Dict[str, Any]:
        """Get statistics about the ChromaDB collection"""
        try:
            count = self.collection.count()
            return {
                "total_documents": count,
                "collection_name": "ragify_documents"
            }
        except Exception as e:
            return {
                "error": str(e),
                "collection_name": "ragify_documents"
            }


# Global document processor instance
document_processor = DocumentProcessor()
